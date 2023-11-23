import os
import sys
import subprocess
from PyQt5.QtWidgets import (QFileDialog,QTableWidget,QTableWidgetItem,QDialog,QApplication, QMainWindow, QPushButton, QMessageBox, QLabel,
                             QVBoxLayout, QWidget, QGridLayout, QGroupBox, QTextEdit,QHeaderView)
from PyQt5.QtGui import QPalette, QColor
from PyQt5.QtCore import Qt
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill

class CustomTitleBar(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAutoFillBackground(True)
        self.setStyleSheet("""
            background-color: #FF17365D;
            color: white;
            font-size: 16px;
            padding: 5px;
        """)
        layout = QVBoxLayout()
        title_label = QLabel("System's Hardware Map")
        layout.addWidget(title_label)
        self.setLayout(layout)
        
class HardwareMapWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.hardware_output = self.get_hardware_info()
        self.initUI()

    def initUI(self):
        self.setWindowTitle("System's Hardware Map")
        # Define a stylesheet for the window's title bar
        title_bar_stylesheet = """
            QMenuBar {
                background-color: #FF17365D; /* Background color of the title bar */
                color: white; /* Title text color */
                font-size: 16px; /* Font size of the title text */
            }
            QMenuBar::item {
                padding: 5px 10px; /* Padding around the title text */
            }
        """

        # Apply the stylesheet to the QMainWindow's menu bar
        self.menuBar().setStyleSheet(title_bar_stylesheet)
        self.setGeometry(100, 100, 1100, 700)
        self.setStyleSheet("background-color: #EEF5FF;")

        layout = QGridLayout()

        # Central label
        self.table_widget = QTableWidget()
        self.table_widget.setColumnCount(2)
        self.table_widget.setHorizontalHeaderLabels(['Attribute', 'Value'])
        # Set a fixed height for the table
       # self.table_widget.setFixedHeight(300)
        #Disable editing for the entire table
        self.table_widget.setEditTriggers(QTableWidget.NoEditTriggers)
        

        # Parse console output and add "System Vendor Name" and "System Model Name" to the table
        console_output = self.get_hardware_info()
        lines = console_output.split('\n')
        for line in lines:
            if 'System Vendor Name:' in line or 'System Model Name:' in line:
                attribute, value = map(str.strip, line.split(':', 1))
                self.table_widget.insertRow(self.table_widget.rowCount())
                self.table_widget.setItem(self.table_widget.rowCount() - 1, 0, QTableWidgetItem(attribute))
                self.table_widget.setItem(self.table_widget.rowCount() - 1, 1, QTableWidgetItem(value))
            if "System's UUID:" in line or 'Host Name:' in line:
                attribute, value = map(str.strip, line.split(':', 1))
                self.table_widget.insertRow(self.table_widget.rowCount())
                self.table_widget.setItem(self.table_widget.rowCount() - 1, 0, QTableWidgetItem(attribute))
                self.table_widget.setItem(self.table_widget.rowCount() - 1, 1, QTableWidgetItem(value))
            if "CPU Vendor Name:" in line or 'CPU Model Name:' in line:
                attribute, value = map(str.strip, line.split(':', 1))
                self.table_widget.insertRow(self.table_widget.rowCount())
                self.table_widget.setItem(self.table_widget.rowCount() - 1, 0, QTableWidgetItem(attribute))
                self.table_widget.setItem(self.table_widget.rowCount() - 1, 1, QTableWidgetItem(value))
            if "CPU Cores:" in line or 'CPU(s):' in line:
                attribute, value = map(str.strip, line.split(':', 1))
                self.table_widget.insertRow(self.table_widget.rowCount())
                self.table_widget.setItem(self.table_widget.rowCount() - 1, 0, QTableWidgetItem(attribute))
                self.table_widget.setItem(self.table_widget.rowCount() - 1, 1, QTableWidgetItem(value))
            if "CPU Architecture:" in line or 'Minimum CPU Frequency:' in line:
                attribute, value = map(str.strip, line.split(':', 1))
                self.table_widget.insertRow(self.table_widget.rowCount())
                self.table_widget.setItem(self.table_widget.rowCount() - 1, 0, QTableWidgetItem(attribute))
                self.table_widget.setItem(self.table_widget.rowCount() - 1, 1, QTableWidgetItem(value))
            if "Maximum CPU Frequency:" in line:
                attribute, value = map(str.strip, line.split(':', 1))
                self.table_widget.insertRow(self.table_widget.rowCount())
                self.table_widget.setItem(self.table_widget.rowCount() - 1, 0, QTableWidgetItem(attribute))
                self.table_widget.setItem(self.table_widget.rowCount() - 1, 1, QTableWidgetItem(value))
            if  'Range Size:' in line:
                attribute, value = map(str.strip, line.split(':', 1))
                self.table_widget.insertRow(self.table_widget.rowCount())
                self.table_widget.setItem(self.table_widget.rowCount() - 1, 0, QTableWidgetItem("RAM"))
                self.table_widget.setItem(self.table_widget.rowCount() - 1, 1, QTableWidgetItem(value))
            if "Disk Size - /home (GiB):" in line:
                attribute, value = map(str.strip, line.split(':', 1))
                self.table_widget.insertRow(self.table_widget.rowCount())
                self.table_widget.setItem(self.table_widget.rowCount() - 1, 0, QTableWidgetItem(attribute))
                self.table_widget.setItem(self.table_widget.rowCount() - 1, 1, QTableWidgetItem(value))
            if 'Resolution:' in line:
                attribute, value = map(str.strip, line.split(':', 1))
                self.table_widget.insertRow(self.table_widget.rowCount())
                self.table_widget.setItem(self.table_widget.rowCount() - 1, 0, QTableWidgetItem("Display Resolution"))
                self.table_widget.setItem(self.table_widget.rowCount() - 1, 1, QTableWidgetItem(value))
            if "GPU Vendor Name:" in line or 'GPU Device Name:' in line:
                attribute, value = map(str.strip, line.split(':', 1))
                self.table_widget.insertRow(self.table_widget.rowCount())
                self.table_widget.setItem(self.table_widget.rowCount() - 1, 0, QTableWidgetItem(attribute))
                self.table_widget.setItem(self.table_widget.rowCount() - 1, 1, QTableWidgetItem(value))

        self.table_widget.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        layout.addWidget(self.table_widget,0,1, 3, layout.columnCount())
        ## self.hardwareLabel = QLabel("     System's Hardware", self)
       ## layout.addWidget(self.hardwareLabel, 1, 1)
        common_stylesheet_QGroupBox_btn= """
            QPushButton {
                background-color: #FF17365D;
                color: rgb(255, 255, 255);
                border-radius: 15px;
                padding: 5px 15px;
            }
        """

        download_btn = QPushButton("Download", self)
        download_btn.clicked.connect(self.saveHardwareInfo)
        download_btn.setStyleSheet(common_stylesheet_QGroupBox_btn)
        layout.addWidget(download_btn,4,2)

         # Compare button
        self.compare_btn = QPushButton("Compare", self)
        self.compare_btn.clicked.connect(self.compareFiles)
        self.compare_btn.setStyleSheet(common_stylesheet_QGroupBox_btn)
        layout.addWidget(self.compare_btn,5,2)

        # Styling
        buttonSize = (100, 50)
        colorPalette = {
            'Power': '#9EB8D9',  # Gold
            'Memory': '#9EB8D9',  # OrangeRed
            'Processing': '#9EB8D9',  # DodgerBlue
            'Storage': '#9EB8D9',  # LimeGreen
            'Network & Connectivity': '#9EB8D9',  
            'I/O Devices': '#9EB8D9'  
        }
       
        color = '#9EB8D9'  # Background color
        font_color = '#374055'  # Font color
        hover_color = '#A25772'  # Background color on hover

        #Helper function to style buttons
        def styleButton(button, category):
            button.setFixedSize(*buttonSize)
            #color = colorPalette.get(category)
            button.setStyleSheet(
                f"""
                QPushButton {{
                    background-color: {color};
                    color: {font_color};
                    border-radius: 5px;
                }}
                QPushButton:hover {{
                    background-color: {hover_color};
                }}
                """
            )
            #button.setStyleSheet(f"background-color: {color}; border-radius: 5px;")

         # Helper function to center widgets in a QGroupBox
        def centeredLayout():
            layout = QVBoxLayout()
            layout.setAlignment(Qt.AlignCenter)
            return layout

        # Memory related components in a group box
        # Define the common stylesheet as a string
        
        common_stylesheet_QGroupBox = """
            QGroupBox {
                border: 1px solid gray;
                border-color: #FF17365D;
                margin-top: 27px;
                font-size: 14px;
                padding: 5px 0px;
                border-bottom-left-radius: 15px;
                border-bottom-right-radius: 15px;
            }

            QGroupBox::title {
                subcontrol-origin: margin;
                subcontrol-position: top center;
                padding: 5px 140px;
                background-color: #FF17365D;
                color: rgb(255, 255, 255);
            }
        """
       
        memoryGroup = QGroupBox("Memory", self)
       # memoryGroup.setFixedSize(300, 200)
        memoryGroup.setStyleSheet(common_stylesheet_QGroupBox)
        memoryLayout = centeredLayout()
        self.ramButton = QPushButton('RAM', self)
        #self.romButton = QPushButton('ROM', self)
        memoryLayout.addWidget(self.ramButton)
        #memoryLayout.addWidget(self.romButton)
        memoryGroup.setLayout(memoryLayout)
        styleButton(self.ramButton, 'Memory')
        #styleButton(self.romButton, 'Memory')
        self.ramButton.clicked.connect(self.showRAMInfo)
        #self.romButton.clicked.connect(self.showRAMInfo)
        layout.addWidget(memoryGroup, 0, 0)

        # Adding another group as an example:
        processingGroup = QGroupBox("Processing", self)
        
        processingGroup.setStyleSheet(common_stylesheet_QGroupBox)
        #processingGroup.setFixedSize(300, 200)
        processingLayout = centeredLayout()
        self.cpuButton = QPushButton('CPU', self)
        self.gpuButton = QPushButton('GPU', self)
        processingLayout.addWidget(self.cpuButton)
        processingLayout.addWidget(self.gpuButton)
        processingGroup.setLayout(processingLayout)
        styleButton(self.cpuButton, 'Processing')
        styleButton(self.gpuButton, 'Processing')
        self.cpuButton.clicked.connect(self.showCPUInfo)
        self.gpuButton.clicked.connect(self.showGPUInfo)
        layout.addWidget(processingGroup, 0, 2)

        # Power category
        powerGroup = QGroupBox("Power", self)
        powerGroup.setStyleSheet(common_stylesheet_QGroupBox)
        powerLayout = centeredLayout()
        self.batteryButton = QPushButton('Battery', self)
        powerLayout.addWidget(self.batteryButton)
        styleButton(self.batteryButton, 'Power')
        powerGroup.setLayout(powerLayout)
        self.batteryButton.clicked.connect(self.showBatteryInfo)
        layout.addWidget(powerGroup, 1, 0)

        # Storage category
        storageGroup = QGroupBox("Storage", self)
        storageGroup.setStyleSheet(common_stylesheet_QGroupBox)
        storageLayout = centeredLayout()
        self.hddButton = QPushButton('HDD/SDD', self)
        storageLayout.addWidget(self.hddButton)
        storageGroup.setLayout(storageLayout)
        styleButton(self.hddButton, 'Storage')
        self.hddButton.clicked.connect(self.showDiskInfo)
        layout.addWidget(storageGroup, 1, 2)

        # Network & Connectivity category
        networkGroup = QGroupBox("Network & Connectivity", self)
        networkGroup.setStyleSheet(common_stylesheet_QGroupBox)
        networkLayout = centeredLayout()
        self.wifiButton = QPushButton('WiFi', self)
        self.ethernetButton = QPushButton('Ethernet', self)
        self.bluetoothButton = QPushButton('Bluetooth', self)
        networkLayout.addWidget(self.wifiButton)
        networkLayout.addWidget(self.ethernetButton)
        networkLayout.addWidget(self.bluetoothButton)
        networkGroup.setLayout(networkLayout)
        styleButton(self.wifiButton, 'Network & Connectivity')
        styleButton(self.ethernetButton, 'Network & Connectivity')
        styleButton(self.bluetoothButton, 'Network & Connectivity')
        self.wifiButton.clicked.connect(self.showWiFiInfo)
        self.ethernetButton.clicked.connect(self.showEthernetInfo)
        self.bluetoothButton.clicked.connect(self.showBluetoothInfo)
        layout.addWidget(networkGroup, 2, 0)

        # I/O Device category
        ioGroup = QGroupBox("I/O Devices", self)
        ioGroup.setStyleSheet(common_stylesheet_QGroupBox)
        ioLayout = centeredLayout()
        self.keyboardButton = QPushButton('Keyboard', self)
        self.touchpadButton = QPushButton('Touchpad', self)
        self.cameraButton = QPushButton('Camera', self)
        self.speakerButton = QPushButton('Speaker', self)
        self.mouseButton = QPushButton('Mouse', self)
        ioLayout.addWidget(self.keyboardButton)
        ioLayout.addWidget(self.touchpadButton)
        ioLayout.addWidget(self.cameraButton)
        ioLayout.addWidget(self.speakerButton)
        ioLayout.addWidget(self.mouseButton)
        ioGroup.setLayout(ioLayout)
        styleButton(self.keyboardButton, 'I/O Devices')
        styleButton(self.touchpadButton, 'I/O Devices')
        styleButton(self.cameraButton, 'I/O Devices')
        styleButton(self.speakerButton, 'I/O Devices')
        styleButton(self.mouseButton, 'I/O Devices')
        self.keyboardButton.clicked.connect(self.showKeyboardInfo)
        self.touchpadButton.clicked.connect(self.showTouchpadInfo)
        self.speakerButton.clicked.connect(self.showSpeakerInfo)
        self.mouseButton.clicked.connect(self.showMouseInfo)
        self.cameraButton.clicked.connect(self.showCameraInfo)
        layout.addWidget(ioGroup, 2, 2)

        # Add central widget to the main window
        centralWidget = QWidget(self)
        centralWidget.setLayout(layout)
        self.setCentralWidget(centralWidget)

        

    
    def get_hardware_info(self):
        # Execute the command and get the output
        command = ["sudo","hardware_map", "all"]
        result = subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)

        # Check for errors
        if result.returncode != 0:
            print("Error:", result.stderr)
            QMessageBox.warning(self, "Error", "Failed to retrieve hardware information.")
            return ""

        return result.stdout

    def extract_info(self, section_header):
        # Extract specific section details from the saved output
        section_start = self.hardware_output.find(f"-------{section_header}")
        if section_start != -1:
            section_info = self.hardware_output[section_start:]
            section_end = section_info.find("-------", 1)  # Find the next "-------" separator
            if section_end != -1:
                section_info = section_info[:section_end]
            return section_info
        else:
            return f"{section_header} details not found in the output."

    def showBatteryInfo(self):
        battery_info = self.extract_info("Battery Details")
        self.showTableDialog("Battery Information", battery_info)

    def showCPUInfo(self):
        cpu_info = self.extract_info("CPU Details")
        self.showTableDialog("CPU Information", cpu_info)

    def showGPUInfo(self):
        cpu_info = self.extract_info("GPU Details")
        self.showTableDialog("GPU Information", cpu_info)
        
    def showDiskInfo(self):
        disk_info = self.extract_info("Disk Details")
        self.showTableDialog("Battery Information", disk_info)
    
    def showRAMInfo(self):
        ram_info = self.extract_info("RAM Details")
        self.showTableDialog("RAM Information", ram_info)

    def showWiFiInfo(self):
        wifi_info = self.extract_info("Wifi Details")
        self.showTableDialog("WiFi Information", wifi_info)   

    def showEthernetInfo(self):
        eth_info = self.extract_info("Ethernet Details")
        self.showTableDialog("Ethernet Information", eth_info)   

    def showBluetoothInfo(self):
        bluetooth_info = self.extract_info("Bluetooth Details")
        self.showTableDialog("Bluetooth Information", bluetooth_info)   

    def showKeyboardInfo(self):
        keyboard_info = self.extract_info("Keyboard Details")
        self.showTableDialog("Keyboard Information", keyboard_info)   

    def showTouchpadInfo(self):
        touchpad_info = self.extract_info("Touchpad Details")
        self.showTableDialog("Touchpad Information", touchpad_info)   

    def showSpeakerInfo(self):
        audio_info = self.extract_info("Audio Details")
        self.showTableDialog("Aduio Information", audio_info)   

    def showMouseInfo(self):
        mouse_info = self.extract_info("Mouse Details")
        self.showTableDialog("Mouse Information", mouse_info)

    def showCameraInfo(self):
        camera_info = self.extract_info("Webcam Details")
        self.showTableDialog("Camera Information", camera_info)
    
    def showMotherboardInfo(self):
        motherboard_info = "Example Motherboard Info: ASUS X570"  # Fetch real info
        QMessageBox.information(self, 'Motherboard Information', motherboard_info)

    # def showBatteryInfo(self):
    #     battery_info = "Example Battery Info: 5000mAh"  # Fetch real info
    #     QMessageBox.information(self, 'Battery Information', battery_info)

   

    # Define other methods for other components similarly...

    def showTableDialog(self, title, info):
        dialog = QDialog(self)
        dialog.setWindowTitle(title)
        dialog.resize(500, 400)

        layout = QVBoxLayout()

        table = QTableWidget()
        table.setColumnCount(2)
        table.setHorizontalHeaderLabels(["Attribute", "Value"])
        table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)

        info_lines = [line.strip() for line in info.splitlines() if ':' in line]
        table.setRowCount(len(info_lines))

        for idx, line in enumerate(info_lines):
            attribute, value = line.split(":", 1)
            table.setItem(idx, 0, QTableWidgetItem(attribute.strip()))
            table.setItem(idx, 1, QTableWidgetItem(value.strip()))

        layout.addWidget(table)
        dialog.setLayout(layout)
        dialog.exec_()
    

    def saveHardwareInfo(self):
        hardware_info = self.get_hardware_info()
        #doc = self.createDoc(hardware_info)

        # Ask the user where to save the doc file
        options = QFileDialog.Options()
        #filePath, _ = QFileDialog.getSaveFileName(self, "Save File", "", "Word Documents (*.docx);;All Files (*)", options=options)
        filePath, _ = QFileDialog.getSaveFileName(self, "Save File", "", "Text Documents (*.txt);;All Files (*)", options=options)

        if filePath:
            if not filePath.endswith('.txt'):
                filePath += '.txt'
            with open(filePath, 'w') as file:
                file.write(hardware_info)

    def compareFiles(self):
        options = QFileDialog.Options()
        filePaths, _ = QFileDialog.getOpenFileNames(self, "Select Files", "", "Text Documents (*.txt);;All Files (*)", options=options)

        if len(filePaths) < 2:
            QMessageBox.warning(self, "Warning", "Please select at least two files to compare.")
            return

        fileNames = [os.path.basename(file) for file in filePaths]
        
        # Process file contents
        processed_contents = []
        for file in filePaths:
            with open(file, 'r') as f:
                lines = f.readlines()
                attributes = []
                values = []
                for line in lines:
                    line = line.strip()  # Remove any leading/trailing whitespace
                    if "-------" in line:
                        attributes.append(line)
                        values.append(None)  # Append None or '' to keep the alignment of attributes and values consistent
                    elif ":" in line:
                        attr, val = line.split(":", 1)
                        attributes.append(attr.strip())
                        values.append(val.strip())
                processed_contents.append((attributes, values))

        # Combine attributes and values for the comparison dialog
        combined_contents = []
        for attributes, values in processed_contents:
            combined_contents.append(["{}: {}".format(attr, val) for attr, val in zip(attributes, values)])
        
        comparisonDialog = ComparisonDialog(combined_contents, fileNames, self)
        comparisonDialog.exec_()
        # if filePath:
        #     if not filePath.endswith('.docx'):
        #         filePath += '.docx'
        #     doc.save(filePath)

    # def createDoc(self, info):
    #     doc = Document()

    #     # Setting the global style for the document
    #     style = doc.styles['Normal']
    #     font = style.font
    #     font.name = 'Arial'  # setting font type
    #     font.size = Pt(12)   # setting font size
    #     font.color.rgb = RGBColor(0, 0, 0)  # setting font color to black

    #     doc.add_heading('Hardware Information', level=1)

    #     sections = info.split("-------")[1:]  # split by "-------" and ignore the first empty part
    #     for i in range(0, len(sections), 2):  # jump in steps of 2 because one part is the header and the other is the content
    #         section_header = sections[i].strip()
    #         section_content = sections[i+1].strip()

    #         doc.add_heading(section_header, level=2)
            
    #         for line in section_content.splitlines():
    #             if ':' in line:
    #                 doc.add_paragraph(line)

    #     return doc

# New Dialog for Comparison

def diff_strings(a, b):
    """Function to compare two strings and highlight the differences."""
    a_label, a_value = a.split(':', 1) if ':' in a else (a, "")
    b_label, b_value = b.split(':', 1) if ':' in b else (b, "")
    
    # Check if the values differ and highlight if necessary
    if a_value != b_value:
        a_value = f'<span style="background-color:yellow;">{a_value}</span>'
        b_value = f'<span style="background-color:yellow;">{b_value}</span>'

    return a_label, a_value.strip(), b_value.strip()

class ComparisonDialog(QDialog):
    def __init__(self, contents, fileNames, parent=None):
        super(ComparisonDialog, self).__init__(parent)
        self.setWindowTitle("File Comparison")
        self.setGeometry(100, 100, 800, 600)
        
        layout = QVBoxLayout(self)
        self.tableWidget = QTableWidget(self)
        self.tableWidget.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOn)
        layout.addWidget(self.tableWidget)

        # Add the download button
        self.downloadButton = QPushButton("Download Comparison", self)
        self.downloadButton.clicked.connect(self.downloadComparison)
        layout.addWidget(self.downloadButton)

        self.compareContents(contents, fileNames)

    
    def downloadComparison(self):
        """Function to download the comparison data as an HTML file."""
        filepath, _ = QFileDialog.getSaveFileName(self, "Save File", "", "HTML Files (*.html);;All Files (*)")
        
        if not filepath:
            return

        if not filepath.endswith('.html'):
            filepath += '.html'

        # Start HTML content with table opening tag
        html_content = """
        <html>
        <head>
            <style>
                table {
                    width: 100%;
                    border-collapse: collapse;
                }
                th, td {
                    border: 1px solid black;
                    padding: 8px;
                    text-align: left;
                }
                tr:hover {background-color: #f5f5f5;}
            </style>
        </head>
        <body>
        <table>
        """

        # Table headers
        headers = ['Attributes']
        for col in range(self.tableWidget.columnCount()):
            item = self.tableWidget.horizontalHeaderItem(col)
            if item:
                headers.append(item.text())

        html_content += '<tr>'
        for header in headers:
            html_content += f'<th>{header}</th>'
        html_content += '</tr>'

        # Table content
        for row in range(self.tableWidget.rowCount()):
            items = [self.tableWidget.item(row, col) for col in range(self.tableWidget.columnCount())]

            html_content += '<tr>'

            # Check if any item in this row contains the "-------" pattern
            if any("-------" in (item.text() if item else "") for item in items):
                html_content += '<tr style="background-color:lightgray;">'  # Change the background color for the entire row
            else:
                html_content += '<tr>'

            # Compare items and highlight differences
            if len(items) > 1 and all(i is not None for i in items):
                label_text, first_val, second_val = diff_strings(items[0].text(), items[1].text())
                # Replace "-------" with an empty string
                label_text = label_text.replace("-------", "")
                first_val = first_val.replace("-------", "")
                second_val = second_val.replace("-------", "")
                html_content += f'<td>{label_text}</td><td>{first_val}</td><td>{second_val}</td>'
            else:
                for item in items:
                    if item:
                        cell_content = item.text().replace("-------", "")  # Replace "-------" with an empty string
                        html_content += f'<td colspan="{len(headers)}">{cell_content}</td>'
                    else:
                        html_content += '<td></td>'
            html_content += '</tr>'

        # Close table and HTML tags
        html_content += """
        </table>
        </body>
        </html>
        """

        # Write to the file
        with open(filepath, 'w') as f:
            f.write(html_content)

        QMessageBox.information(self, "Success", "File has been saved successfully!")

    def compareContents(self, contents, fileNames):
        # Number of rows is the max number of lines in any file, number of columns is 1 + the number of files (for attribute names)
        numRows = max(len(content) for content in contents)
        self.tableWidget.setRowCount(numRows)
        self.tableWidget.setColumnCount(len(fileNames) + 1) # +1 for the attribute names

        headers = ["Attributes"] + fileNames
        self.tableWidget.setHorizontalHeaderLabels(headers)

         # Fill in attribute names and values
        for rowIndex, lines in enumerate(zip(*contents)): # Transpose to get rows
            attribute = lines[0].split(":")[0] if ":" in lines[0] else lines[0]
            attribute = attribute.replace("-------", "").strip()  # Remove "-------"
            
            # Create an item for the attribute
            attributeItem = QTableWidgetItem(attribute)
            isHeader = "-------" in lines[0]  # This will check if original line had "-------"
            
            if isHeader:  # If this row corresponds to a section header
                attributeItem.setBackground(QColor(200, 200, 200))  # Set a gray background color
            self.tableWidget.setItem(rowIndex, 0, attributeItem)
            
            for columnIndex, line in enumerate(lines, start=1): # Start from 1 to skip attribute column
                value = line.split(":")[1].strip() if ":" in line else ""
                if value == "None":
                    value = ""  # Set it as empty if the value is "None"
                item = QTableWidgetItem(value)
                
                if isHeader:  # If this row corresponds to a section header
                    item.setBackground(QColor(200, 200, 200))  # Set a gray background color
                self.tableWidget.setItem(rowIndex, columnIndex, item)

        # Highlight differences
        self.highlightDifferences(contents)

    def highlightDifferences(self, contents):
        for i in range(self.tableWidget.rowCount()):
            values = [self.tableWidget.item(i, j).text() for j in range(1, self.tableWidget.columnCount())] # Skip attribute column
            if len(set(values)) > 1: # More than 1 unique value implies differences
                for j in range(1, self.tableWidget.columnCount()): # Again, skip attribute column
                    self.tableWidget.item(i, j).setBackground(QColor(255, 255, 0))  # Yellow background for differences

    # def compareContents(self, contents, fileNames):
    #     max_rows = max([len(content) for content in contents])
    #     column_count = len(contents)
        
    #     self.tableWidget.setRowCount(max_rows)
    #     self.tableWidget.setColumnCount(column_count)
    #     self.tableWidget.setHorizontalHeaderLabels(fileNames)
        
    #     for col_num, content in enumerate(contents):
    #         for row_num, line in enumerate(content):
    #             item = QTableWidgetItem(line)
    #             item.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)  # Make cell non-editable
    #             self.tableWidget.setItem(row_num, col_num, item)
        
    #     # Check differences
    #     for row in range(max_rows):
    #         unique_values = set()
    #         for col in range(column_count):
    #             item = self.tableWidget.item(row, col)
    #             if item:
    #                 unique_values.add(item.text().strip())
    #         if len(unique_values) > 1:  # meaning there are differences
    #             for col in range(column_count):
    #                 item = self.tableWidget.item(row, col)
    #                 if item:
    #                     item.setBackground(QColor(255, 0, 0))

    #     self.tableWidget.horizontalHeader().setSectionResizeMode(QHeaderView.Interactive)


if __name__ == '__main__':
    app = QApplication(sys.argv)
    mainWindow = HardwareMapWindow()
    mainWindow.show()
    sys.exit(app.exec_())

